"""
Extraccion de texto de documentos Word y PDF.
Usa MarkItDown como primario, pymupdf como fallback para PDFs.
"""

import sys
from pathlib import Path
from config import logger

def extract_text(file_path: str) -> str:
    """
    Extrae todo el texto de un archivo .docx o .pdf.
    Retorna string con el texto extraido.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".docx":
        return _extract_docx(path)
    elif suffix == ".pdf":
        return _extract_pdf(path)
    else:
        logger.warning(f"Formato no soportado: {suffix} — {path.name}")
        return ""


def _extract_docx(path: Path) -> str:
    """Extrae texto de .docx usando python-docx (ligero y rapido)."""
    try:
        from docx import Document
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Extraer tambien texto de tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except ImportError:
        # Fallback a MarkItDown si python-docx no esta
        return _extract_with_markitdown(path)
    except Exception as e:
        logger.error(f"Error extrayendo {path.name}: {e}")
        return ""


def _extract_pdf(path: Path) -> str:
    """Extrae texto de PDF. Primero pymupdf (rapido), fallback MarkItDown."""
    # Intento 1: pymupdf (mas rapido, nativo)
    try:
        import fitz
        doc = fitz.open(str(path))
        pages = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                pages.append(text.strip())
        doc.close()
        if pages:
            return "\n\n".join(pages)
    except ImportError:
        pass
    except Exception as e:
        logger.debug(f"pymupdf fallo con {path.name}: {e}")

    # Intento 2: MarkItDown
    return _extract_with_markitdown(path)


def _extract_with_markitdown(path: Path) -> str:
    """Extrae texto usando MarkItDown (maneja docx, pdf, pptx, etc.)."""
    try:
        from markitdown import MarkItDown
        md = MarkItDown()
        result = md.convert(str(path))
        return result.text_content
    except ImportError:
        logger.error("MarkItDown no instalado. Ejecuta: pip install markitdown")
        return ""
    except Exception as e:
        logger.error(f"MarkItDown fallo con {path.name}: {e}")
        return ""


# ---------------------------------------------------------------------------
# Discovery — encontrar todos los archivos procesables en una carpeta
# ---------------------------------------------------------------------------

def discover_files(folder: str, extensions=None) -> list:
    """
    Encuentra todos los archivos .docx y .pdf en una carpeta (recursivo).
    Retorna lista de Path objects.
    """
    if extensions is None:
        extensions = {".docx", ".pdf"}

    folder = Path(folder)
    if not folder.exists():
        logger.error(f"Carpeta no encontrada: {folder}")
        return []

    files = []
    for ext in extensions:
        files.extend(folder.rglob(f"*{ext}"))

    logger.info(f"Encontrados {len(files)} archivos en {folder}")
    return sorted(files)
